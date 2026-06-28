"""
extractor.py
-------------
Handles getting raw text out of uploaded files:
- Resume: PDF (primary), DOCX, TXT
- Job Description: PDF, DOCX, TXT, or pasted text

Uses pdfplumber as the primary PDF engine (better layout handling)
and falls back to PyPDF2 if pdfplumber fails on a malformed file.
"""

import io
import re

import pdfplumber
import PyPDF2
from docx import Document


def _clean_text(text: str) -> str:
    """Normalize whitespace, fix common PDF extraction artifacts."""
    if not text:
        return ""
    # Fix hyphenated line breaks: "develop-\nment" -> "development"
    text = re.sub(r"-\n", "", text)
    # Collapse multiple newlines/spaces but keep paragraph structure
    text = re.sub(r"\r\n", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using pdfplumber, falling back to PyPDF2."""
    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_parts.append(page_text)
        text = "\n".join(text_parts)
        if text.strip():
            return _clean_text(text)
    except Exception:
        pass

    # Fallback: PyPDF2
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        text_parts = [page.extract_text() or "" for page in reader.pages]
        return _clean_text("\n".join(text_parts))
    except Exception as e:
        raise ValueError(f"Could not extract text from PDF: {e}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a .docx file's bytes."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        return _clean_text("\n".join(parts))
    except Exception as e:
        raise ValueError(f"Could not extract text from DOCX: {e}")


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Decode a plain text file, trying common encodings."""
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return _clean_text(file_bytes.decode(encoding))
        except (UnicodeDecodeError, AttributeError):
            continue
    return _clean_text(file_bytes.decode("utf-8", errors="ignore"))


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Dispatch extraction based on file extension."""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext == "docx":
        return extract_text_from_docx(file_bytes)
    elif ext in ("txt", "md"):
        return extract_text_from_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: .{ext}. Please upload PDF, DOCX, or TXT.")


def get_page_count(file_bytes: bytes, filename: str) -> int:
    """Return page count for a PDF (used in ATS heuristics); 1 for other types."""
    if not filename.lower().endswith(".pdf"):
        return 1
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            return len(pdf.pages)
    except Exception:
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            return len(reader.pages)
        except Exception:
            return 1
